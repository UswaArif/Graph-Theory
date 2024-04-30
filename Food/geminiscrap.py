from bs4 import BeautifulSoup
import pandas as pd
from selenium import webdriver
from docx import Document

def scrape_data(driver):
    urlList = [
        "https://www.readyseteat.com/inspiration-and-tips/earn-these-graduation-party-food-ideas",
        "https://www.readyseteat.com/inspiration-and-tips/dig-delicious-dips-gardeins-ultimate-plant-based-chickn",
        "https://www.readyseteat.com/inspiration-and-tips/grocery-budget-tips-and-tricks",
        "https://www.readyseteat.com/inspiration-and-tips/leftover-ham-recipes",
        "https://www.readyseteat.com/inspiration-and-tips/makerz-bakerz-easy-frosting-hacks",
        "https://www.readyseteat.com/inspiration-and-tips/12-meatless-monday-ideas-start-week-right",
        "https://www.readyseteat.com/inspiration-and-tips/vegan-easter-recipes",
        "https://www.readyseteat.com/inspiration-and-tips/copycat-restaurant-recipes",
        "https://www.readyseteat.com/inspiration-and-tips/get-fired-all-grill-fathers-day-menu",
        "https://www.readyseteat.com/inspiration-and-tips/how-to-make-grain-bowls",
        "https://www.readyseteat.com/inspiration-and-tips/festive-cinco-de-mayo-recipes",
        "https://www.readyseteat.com/inspiration-and-tips/easter-dessert-dinner-recipes",
        "https://www.readyseteat.com/inspiration-and-tips/vegan-easter-recipes",
        "https://www.readyseteat.com/inspiration-and-tips/ultimate-spring-produce-recipe-guide",
        "https://www.readyseteat.com/inspiration-and-tips/chinese-new-year-tradition-easy-dumpling-recipes"
    ]

    count = 1
    for url in urlList:
        try:
            driver.get(url)
            content = driver.page_source
            soup = BeautifulSoup(content, "html5lib")
            document = Document()
            #document.add_heading(f'Scraped Food Data - {url}', level=1)
            for name in soup.findAll('div', {'class': 'c-article-content__body lazyloaded'}):
                name2 = name.find('div', {'class': 'c-article-section c-article-section__introduction'})
                if name2 is not None:
                    document.add_paragraph(name2.text)
                else:
                    document.add_paragraph(" ")
            document_path = f'D:\\6 semester\\GT project\\fooddata\\scraped_data_{count}.docx'
            document.save(document_path)
            count += 1
           
        except Exception as e:
            print(f"Error scraping {url}: {e}")

    driver.quit()

options = webdriver.ChromeOptions()
options.add_argument('executable_path=c:\\Users\\dell\\Downloads\\chromedriver.exe')  # Replace with your path
driver = webdriver.Chrome(options=options)

scrape_data(driver)
