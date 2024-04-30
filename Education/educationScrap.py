from bs4 import BeautifulSoup
import pandas as pd
from selenium import webdriver
from docx import Document

def scrape_data(driver):
    url = 'https://theconversation.com/global/topics/education-274?page='
    count = 1
    for i in range(1,16):
        try:
            i = str(i)
            driver.get(url + i)
            content = driver.page_source
            soup = BeautifulSoup(content, "html5lib")
            document = Document()
            #document.add_heading(f'Scraped Food Data - {url}', level=1)
            for name in soup.findAll('article', {'class': 'clearfix placed analysis published'}):
                name2 = name.find('div', {'class': 'article--header'})
                name3 = name.find('div', {'class': 'content'})
                if name2 is not None:
                    document.add_paragraph(name2.text)
                else:
                    document.add_paragraph(" ")
                if name3 is not None:
                    document.add_paragraph(name3.text)
                else:
                    document.add_paragraph(" ")
            document_path = f'D:\\6 semester\\GT project\\educationdata\\scraped_data_{count}.docx'
            document.save(document_path)
            count += 1
           
        except Exception as e:
            print(f"Error scraping {url}: {e}")

    driver.quit()

options = webdriver.ChromeOptions()
options.add_argument('executable_path=c:\\Users\\dell\\Downloads\\chromedriver.exe')  # Replace with your path
driver = webdriver.Chrome(options=options)

scrape_data(driver)
